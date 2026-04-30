"""
Генерирует список кредиторов, редактируя копию шаблона Елены.
Структура таблицы (doc.tables[1]):
  Rows 0-3:  заголовки Раздела I
  Rows 4-18: данные Раздела I (1.1–1.15)
  Row  19:   заголовок Раздела II
  Row  20:   заголовок колонок Раздела II
  Rows 21-80: данные Раздела II (2.1–2.60)
"""
import json
import shutil
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt


def write_cell(cell, text: str, font_size: int = 9):
    """Очищает ячейку полностью и записывает текст. \n = новый параграф."""
    from lxml import etree
    from docx.oxml.ns import qn as _qn

    tc = cell._tc
    # Удаляем все параграфы кроме первого
    paras = tc.findall(_qn('w:p'))
    for p in paras[1:]:
        tc.remove(p)

    # Очищаем первый параграф — удаляем все runs
    first_p = paras[0] if paras else tc.add_p()
    for r in first_p.findall(_qn('w:r')):
        first_p.remove(r)

    text_str = str(text) if text is not None else ''
    lines = text_str.split('\n')

    # Первая строка в очищенный первый параграф
    from docx.oxml import OxmlElement
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(_qn('w:val'), str(font_size * 2))
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement('w:t')
    t.text = lines[0]
    r.append(t)
    first_p.append(r)

    # Остальные строки — новые параграфы
    for line in lines[1:]:
        new_p = OxmlElement('w:p')
        new_r = OxmlElement('w:r')
        new_rpr = OxmlElement('w:rPr')
        new_sz = OxmlElement('w:sz')
        new_sz.set(_qn('w:val'), str(font_size * 2))
        new_rpr.append(new_sz)
        new_r.append(new_rpr)
        new_t = OxmlElement('w:t')
        new_t.text = line
        new_r.append(new_t)
        new_p.append(new_r)
        tc.append(new_p)


def insert_row_after(table, after_row_idx: int) -> object:
    """Вставляет новую строку после указанной, клонируя её XML."""
    source_tr = table.rows[after_row_idx]._tr
    new_tr = deepcopy(source_tr)
    source_tr.addnext(new_tr)
    # Найти новую строку в таблице
    for row in table.rows:
        if row._tr is new_tr:
            return row
    return table.rows[after_row_idx + 1]


def clear_row_data(row, start_col: int = 0):
    """Очищает все ячейки строки."""
    seen = set()
    for cell in row.cells:
        if id(cell._tc) not in seen:
            seen.add(id(cell._tc))
            write_cell(cell, '')


def fill_section1_row(row, idx: int, creditor: dict):
    """
    Заполняет строку Раздела I.
    Колонки: [0]=№, [2]=содержание, [3]=кредитор, [4]=адрес, [5]=основание, [6]=сумма, [7]=долг, [8]=долг2, [9]=штрафы, [10]=штрафы2
    """
    cells = row.cells
    name = creditor.get('creditor_name', '')
    inn = creditor.get('creditor_inn', '')
    address = creditor.get('address', '')
    debt = creditor.get('current_debt', 0) or 0
    # Если НБКИ-остаток < 100 (технический нуль/остаток после переуступки) — берём сумму из ИП
    if debt < 100 and creditor.get('proceedings'):
        ip_amount = creditor['proceedings'][0].get('amount_principal') or creditor['proceedings'][0].get('amount_total') or 0
        if ip_amount > debt:
            debt = ip_amount

    uid = creditor.get('uid', '')
    contract = creditor.get('contract_number', '')
    raw_date = creditor.get('contract_date', '')
    # Конвертируем дату из DD-MM-YYYY в DD.MM.YYYY
    date_fmt = raw_date.replace('-', '.') if raw_date else ''

    lines = []
    if uid:
        lines.append(f"УИд договора: {uid}")
    if contract and date_fmt:
        lines.append(f"Номер договора: {contract} от {date_fmt}")
    elif contract:
        lines.append(f"Номер договора: {contract}")

    # Добавляем данные из сматченного ИП
    if creditor.get('proceedings'):
        for proc in creditor['proceedings'][:1]:
            if proc.get('basis_doc'):
                lines.append(proc['basis_doc'])
            if proc.get('ip_number'):
                ip_line = f"ИП № {proc['ip_number']}"
                if proc.get('ip_date'):
                    ip_line += f" от {proc['ip_date']}"
                lines.append(ip_line)

    basis = '\n'.join(lines)

    seen = set()
    # [0] и [1] — слитые ячейки, пишем только в [0]
    write_cell(cells[0], f"1.{idx}")
    seen.add(id(cells[0]._tc))
    seen.add(id(cells[1]._tc))

    col_data = {2: 'кредит', 3: f"{name}, ИНН {inn}", 4: address, 5: basis,
                6: f"{debt:.2f}", 7: '00.00', 8: '00.00', 9: '00.00', 10: '00.00'}
    for col_idx, val in col_data.items():
        if col_idx < len(cells) and id(cells[col_idx]._tc) not in seen:
            write_cell(cells[col_idx], val)
            seen.add(id(cells[col_idx]._tc))


def fill_section2_row(row, idx: int, item: dict):
    """
    Заполняет строку Раздела II.
    Уникальные ячейки: [0]=№, [2]=наименование, [8]=недоимка, [9]=штрафы
    """
    cells = row.cells
    amount = item.get('amount', 0) or 0

    seen = set()
    # Запишем только в уникальные ячейки
    col_data = {0: f"2.{idx}", 2: item.get('description', ''), 8: f"{amount:.2f}", 9: '00.00'}
    for col_idx, val in col_data.items():
        if col_idx < len(cells) and id(cells[col_idx]._tc) not in seen:
            write_cell(cells[col_idx], val)
            seen.add(id(cells[col_idx]._tc))


def build_document(matched_data: dict, template_path: str, output_path: str):
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)
    table = doc.tables[1]

    section1 = matched_data['section1']
    section2 = matched_data['section2']

    # ---- Раздел I ----
    # Существующие строки данных: rows 4-18 (15 строк)
    SECTION1_DATA_START = 4
    SECTION1_DATA_END = 18   # включительно
    existing_s1_rows = SECTION1_DATA_END - SECTION1_DATA_START + 1  # = 15

    for i, cred in enumerate(section1):
        row_idx = SECTION1_DATA_START + i
        if row_idx <= SECTION1_DATA_END:
            fill_section1_row(table.rows[row_idx], i + 1, cred)
        else:
            # Нужна новая строка — вставляем перед Разделом II
            new_row = insert_row_after(table, row_idx - 1)
            fill_section1_row(new_row, i + 1, cred)

    # Если кредиторов меньше 15 — очищаем лишние строки
    for i in range(len(section1), existing_s1_rows):
        row_idx = SECTION1_DATA_START + i
        if row_idx < len(table.rows):
            clear_row_data(table.rows[row_idx])

    print(f"Раздел I: {len(section1)} кредиторов записано")

    # ---- Раздел II ----
    # После возможного добавления строк в Раздел I, находим заголовок Раздела II заново
    section2_header_idx = None
    for i, row in enumerate(table.rows):
        if 'Обязательные платежи' in row.cells[2].text:
            section2_header_idx = i
            break

    if section2_header_idx is None:
        print("ПРЕДУПРЕЖДЕНИЕ: не найден заголовок Раздела II")
        doc.save(output_path)
        return

    SECTION2_DATA_START = section2_header_idx + 2   # +1 заголовок колонок, +1 первая строка данных
    existing_s2_rows = len(table.rows) - SECTION2_DATA_START

    for i, item in enumerate(section2):
        row_idx = SECTION2_DATA_START + i
        if row_idx < len(table.rows):
            fill_section2_row(table.rows[row_idx], i + 1, item)
        else:
            new_row = insert_row_after(table, row_idx - 1)
            fill_section2_row(new_row, i + 1, item)

    # Очищаем лишние строки Раздела II
    for i in range(len(section2), existing_s2_rows):
        row_idx = SECTION2_DATA_START + i
        if row_idx < len(table.rows):
            clear_row_data(table.rows[row_idx])
            # Очищаем номер тоже
            write_cell(table.rows[row_idx].cells[0], '')

    print(f"Раздел II: {len(section2)} записей записано")

    doc.save(output_path)
    print(f"\nДокумент сохранён: {output_path}")


def main():
    base = Path(__file__).parent.parent
    with open(base / 'output' / 'matched.json', encoding='utf-8') as f:
        matched = json.load(f)

    build_document(
        matched,
        str(base / 'input' / 'creditors_template.docx'),
        str(base / 'output' / 'final_creditors_list.docx'),
    )


if __name__ == '__main__':
    main()
